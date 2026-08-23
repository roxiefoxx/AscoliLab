from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "docs"
ASSET_DIR = OUTPUT_DIR / "assets"
OUTPUT_DOCX = OUTPUT_DIR / "ee_backbone_method_brief_for_advisor.docx"
METHOD_SUMMARY = ROOT / "ee_backbone_comparison_outputs" / "method_summary.csv"
LONG_RESULTS = ROOT / "ee_backbone_comparison_outputs" / "method_comparison_long.csv"
MATRIX_PATH = ROOT / "matrices" / "mij_EE_matrix.csv"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
RULE = "D9E2F3"


METHOD_LABELS = {
    "milp_subtour_elimination": "MILP, maximum simple path",
    "maximum_weight_asymmetric_hamiltonian_path": "Hamiltonian path (all nodes)",
    "branch_and_bound": "Branch-and-bound",
    "greedy_tree_path": "Greedy tree/path",
    "dynamic_programming": "Dynamic programming / beam DP",
    "maximum_spanning_tree": "Maximum spanning tree path",
}


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths: Iterable[float]):
    widths = list(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    table._tbl.insert(0, grid)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_bottom_border(paragraph, color="D9E2F3", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_paragraph(doc, text="", style=None, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic, color=color or INK)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, color=INK)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "E-to-E Backbone Method Brief"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_run_font(run, size=9, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Prepared from ee_backbone_method_comparison.ipynb | August 20, 2026"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        set_run_font(run, size=9, color=GRAY)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("METHOD BRIEF")
    set_run_font(r, size=10, bold=True, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("E-to-E Backbone Method Comparison")
    set_run_font(r, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(
        "Advisor submission summary for the notebook, data scope, optimization methods, and current results"
    )
    set_run_font(r, size=12, color=GRAY)

    metadata = [
        ("Prepared for", "Advisor review"),
        ("Prepared from", "ee_backbone_method_comparison.ipynb"),
        ("Primary input", "matrices/mij_EE_matrix.csv"),
        ("Date", "August 20, 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(value)
        set_run_font(r, color=INK)

    rule = doc.add_paragraph()
    add_bottom_border(rule, color=RULE, size="8")


def create_bar_chart(summary: pd.DataFrame, output_path: Path):
    ordered = summary.copy()
    ordered["label"] = ordered["method"].map(METHOD_LABELS).fillna(ordered["method"])
    ordered = ordered.sort_values("mean_summed_weight", ascending=True)
    width, height = 1200, 560
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 24) if font_path.exists() else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 20) if font_path.exists() else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 26) if bold_path.exists() else font

    draw.text((40, 24), "Mean summed path weight by method", fill=(0, 0, 0), font=bold)
    max_val = float(ordered["mean_summed_weight"].max())
    x0, x1 = 420, 1090
    y = 92
    bar_h = 38
    gap = 34
    colors = {
        "milp_subtour_elimination": (46, 116, 181),
        "maximum_weight_asymmetric_hamiltonian_path": (91, 155, 213),
        "branch_and_bound": (112, 173, 71),
        "greedy_tree_path": (165, 165, 165),
        "dynamic_programming": (237, 125, 49),
        "maximum_spanning_tree": (127, 127, 127),
    }
    for _, row in ordered.iterrows():
        method = row["method"]
        value = float(row["mean_summed_weight"])
        label = row["label"]
        draw.text((40, y + 4), label, fill=(0, 0, 0), font=small)
        draw.rectangle((x0, y, x1, y + bar_h), fill=(238, 242, 247), outline=(220, 225, 232))
        bar_w = int((x1 - x0) * value / max_val)
        draw.rectangle((x0, y, x0 + bar_w, y + bar_h), fill=colors.get(method, (46, 116, 181)))
        draw.text((x0 + bar_w + 10, y + 5), f"{value:,.0f}", fill=(0, 0, 0), font=small)
        y += bar_h + gap
    output_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(output_path)


def add_summary_table(doc, summary: pd.DataFrame):
    order = [
        "milp_subtour_elimination",
        "maximum_weight_asymmetric_hamiltonian_path",
        "branch_and_bound",
        "greedy_tree_path",
        "dynamic_programming",
        "maximum_spanning_tree",
    ]
    display = summary.set_index("method").loc[order].reset_index()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Method", "Mean weight", "Median weight", "Mean edges", "Max weight"]
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], LIGHT_FILL)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True, color=INK)
    for _, row in display.iterrows():
        cells = table.add_row().cells
        values = [
            METHOD_LABELS[row["method"]],
            f"{row['mean_summed_weight']:,.0f}",
            f"{row['median_summed_weight']:,.0f}",
            f"{row['mean_length']:.1f}",
            f"{row['max_summed_weight']:,.0f}",
        ]
        for i, value in enumerate(values):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)
    set_table_fixed_width(table, [2.45, 1.0, 1.0, 0.9, 1.15])


def add_method_overview_table(doc):
    rows = [
        (
            "Greedy tree/path",
            "Fast local baseline.",
            "Shows what is recovered by repeatedly taking the strongest available outgoing edge.",
        ),
        (
            "Maximum spanning tree path",
            "Sparse structural skeleton.",
            "Collapses the graph into high-weight non-cycling undirected structure, then scores directed paths through it.",
        ),
        (
            "Branch-and-bound",
            "Bounded exhaustive search.",
            "Prunes partial paths whose optimistic upper bound cannot beat the current best path.",
        ),
        (
            "Dynamic programming / beam DP",
            "Subset-state search or approximation.",
            "Exact only when the node set is small enough; beam mode keeps the strongest states for tractability.",
        ),
        (
            "MILP, maximum simple path",
            "Exact maximizing benchmark.",
            "Uses binary edge/node variables and MTZ order constraints to optimize a single simple path from each seed.",
        ),
        (
            "Maximum-weight asymmetric Hamiltonian path",
            "Complete directed ordering.",
            "Forces all 32 excitatory nodes to appear exactly once, preserving directed edge asymmetry.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Method", "Why included", "How to interpret it"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        set_cell_shading(table.rows[0].cells[i], LIGHT_FILL)
    for method, why, interp in rows:
        cells = table.add_row().cells
        cells[0].text = method
        cells[1].text = why
        cells[2].text = interp
    set_table_fixed_width(table, [1.8, 1.7, 3.0])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9.5, bold=(row is table.rows[0]))


def build_document():
    OUTPUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)

    summary = pd.read_csv(METHOD_SUMMARY)
    long = pd.read_csv(LONG_RESULTS)
    matrix = pd.read_csv(MATRIX_PATH, index_col=0).astype(float)
    positive_off_diag = matrix.clip(lower=0).to_numpy()
    for i in range(positive_off_diag.shape[0]):
        positive_off_diag[i, i] = 0.0
    diag = pd.Series([matrix.loc[node, node] for node in matrix.index], index=matrix.index)

    chart_path = ASSET_DIR / "method_mean_weights.png"
    create_bar_chart(summary, chart_path)

    doc = Document()
    set_document_styles(doc)
    add_header_footer(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "Bottom line",
        "The notebook compares six approaches for extracting excitatory-to-excitatory backbone paths from the E-only matrix. The unconstrained MILP with subtour elimination is the exact maximizing benchmark for the current objective: all 32 seed solves returned Optimal, and it produced the highest mean summed path weight. The Hamiltonian MILP also solved optimally for all seeds, but it answers a stricter complete-ordering question because every excitatory node must be visited exactly once.",
    )

    doc.add_heading("1. Purpose of the notebook", level=1)
    add_paragraph(
        doc,
        "The notebook evaluates alternative ways to build directed E-to-E backbone paths from the matrix representation of connectivity weights. Each method starts from every excitatory seed node and produces a directed path scored by the sum of selected positive transition weights between distinct excitatory nodes. The purpose is methodological comparison: which search formulation best captures high-weight directed structure under the same preprocessing and scoring rules?",
    )
    add_paragraph(
        doc,
        "The current analysis uses `mij_EE_matrix.csv` as the input. That file contains the 32 excitatory rows and columns extracted from `mij_matrix.csv`. The notebook treats all row and column labels in this restricted matrix as excitatory nodes, clips negative weights out of the candidate transition graph, and sets the diagonal to zero before path search.",
    )

    doc.add_heading("2. Data scope and graph assumptions", level=1)
    facts = [
        ("Input matrix", "matrices/mij_EE_matrix.csv"),
        ("Nodes", f"{matrix.shape[0]} excitatory nodes"),
        ("Candidate directed E-to-E edges", f"{int((positive_off_diag > 1e-12).sum())} positive off-diagonal transitions"),
        ("Self-connections", f"{int((diag.abs() > 1e-12).sum())} nonzero diagonal entries; {int((diag > 1e-12).sum())} are positive"),
        ("Diagonal handling", "Excluded from all current path-scoring methods"),
        ("Graph structure", "Directed and cyclic; not a DAG"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Current setting"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_FILL)
    for k, v in facts:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    set_table_fixed_width(table, [1.75, 4.75])

    add_paragraph(
        doc,
        "The cyclic structure is important. A DAG longest-path algorithm would not apply, and allowing repeated nodes would make a positive cycle reusable. For that reason, the exact and heuristic path methods enforce simple paths: once a node is used, it cannot be revisited.",
    )

    doc.add_heading("3. Methods compared", level=1)
    add_paragraph(
        doc,
        "The methods are intentionally diverse. Some are fast heuristics, some are bounded searches, and the MILP variants are solver-backed formulations that can certify optimality when the solver status is Optimal.",
    )
    add_method_overview_table(doc)

    doc.add_heading("4. Current results", level=1)
    add_paragraph(
        doc,
        "Both solver-backed methods completed successfully: the unconstrained MILP returned `Optimal` for all 32 seed nodes, and the Hamiltonian MILP returned `Optimal` for all 32 seed nodes. Hamiltonian paths contain 31 edges for every seed, as expected for a 32-node path.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart = p.add_run().add_picture(str(chart_path), width=Inches(6.3))
    chart._inline.docPr.set("title", "Mean summed path weight by method")
    chart._inline.docPr.set(
        "descr",
        "Horizontal bar chart showing the unconstrained MILP has the highest mean summed path weight, followed by the Hamiltonian path, branch-and-bound, greedy, dynamic programming, and maximum spanning tree methods.",
    )
    cap = doc.add_paragraph("Figure 1. Mean summed path weight by method.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        set_run_font(r, size=9, italic=True, color=GRAY)

    add_summary_table(doc, summary)

    doc.add_heading("5. Interpretation", level=1)
    add_paragraph(
        doc,
        "The unconstrained MILP is the strongest benchmark for the stated objective because it maximizes summed path weight over simple directed paths while preventing subtours. Its higher mean score indicates that the earlier heuristic and bounded methods often stop short of the best available combination of directed transitions.",
    )
    add_paragraph(
        doc,
        "The Hamiltonian path is useful for a different scientific question: it gives a complete ordering through all excitatory nodes. Its mean path weight is lower than the unconstrained MILP because it is forced to include every node, even when a complete traversal requires weaker connecting edges. That difference should be interpreted as the cost of complete coverage, not as solver failure.",
    )
    add_paragraph(
        doc,
        "Dynamic programming is retained as a useful conceptual comparison, but at 32 nodes the notebook uses beam-limited DP rather than full bitmask DP. Beam pruning explains why dynamic programming can appear weak in the current results: it is not proving the global maximum at this graph size.",
    )

    doc.add_heading("6. Caveats and recommended next checks", level=1)
    caveats = [
        "The current objective uses positive off-diagonal E-to-E transition weights only. Negative transitions and diagonal self-connections are excluded from path scoring.",
        "Because 27 E-only nodes have positive diagonal entries, a separate self-connection sensitivity analysis may be warranted if self-effects are biologically meaningful.",
        "The unconstrained MILP and Hamiltonian MILP are exact only for rows whose solver status is Optimal. In the current run, all relevant rows are Optimal.",
        "If the advisor wants a complete ordered sequence of all excitatory nodes, use the Hamiltonian path output. If the advisor wants the highest-weight simple backbone of any length, use the unconstrained MILP output.",
    ]
    for item in caveats:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)

    doc.add_heading("7. Output files", level=1)
    outputs = [
        ("Notebook", "ee_backbone_method_comparison.ipynb"),
        ("Input matrix", "matrices/mij_EE_matrix.csv"),
        ("Long comparison table", "ee_backbone_comparison_outputs/method_comparison_long.csv"),
        ("Wide comparison table", "ee_backbone_comparison_outputs/method_comparison_wide.csv"),
        ("Selected path edges", "ee_backbone_comparison_outputs/all_methods_edges.csv"),
        ("Method summary", "ee_backbone_comparison_outputs/method_summary.csv"),
        ("MILP benchmark gaps", "ee_backbone_comparison_outputs/milp_benchmark_gaps.csv"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Artifact"
    table.rows[0].cells[1].text = "Path"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_FILL)
    for artifact, path_text in outputs:
        cells = table.add_row().cells
        cells[0].text = artifact
        cells[1].text = path_text
    set_table_fixed_width(table, [1.8, 4.7])

    doc.add_heading("References", level=1)
    refs = [
        "Held, M., & Karp, R. M. (1962). A dynamic programming approach to sequencing problems. Journal of the Society for Industrial and Applied Mathematics, 10(1), 196-210. https://doi.org/10.1137/0110015.",
        "Miller, C. E., Tucker, A. W., & Zemlin, R. A. (1960). Integer Programming Formulation of Traveling Salesman Problems. Journal of the ACM, 7(4), 326-329. https://doi.org/10.1145/321043.321046.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(ref)
        set_run_font(r, size=10)

    doc.core_properties.title = "E-to-E Backbone Method Comparison"
    doc.core_properties.subject = "Advisor method brief"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_document())
